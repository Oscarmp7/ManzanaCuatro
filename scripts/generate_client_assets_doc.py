from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT_PATH = Path(r"C:\Users\omato\OneDrive\Desktop\Listado_Assets_Cliente_Manzana_Cuatro.docx")


def set_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in (("Title", 20), ("Heading 1", 14), ("Heading 2", 11.5)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Listado de assets multimedia solicitados al cliente")
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proyecto: Manzana Cuatro").bold = True
    subtitle.add_run("\nDocumento para solicitar el material pendiente y poder cerrar la web.")


def add_section(document: Document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=1)
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered_section(document: Document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=1)
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    set_base_styles(document)
    add_title(document)

    intro = document.add_paragraph()
    intro.add_run("Revision del proyecto actual. ").bold = True
    intro.add_run(
        "La web hoy solo tiene 5 imagenes de portfolio reales y 1 imagen de studio. "
        "Los videos de la home y de colorizacion estan en placeholder, asi que ese material "
        "todavia hay que pedirlo al cliente."
    )

    add_section(
        document,
        "1. Resumen ejecutivo",
        [
            "La home usa 4 piezas visuales tipo reel para abrir la experiencia.",
            "La seccion de colorizacion actualmente depende de 3 casos con video o imagen.",
            "La pagina Studio usa una sola imagen y conviene reemplazarla por una foto de equipo grande y profesional.",
            "La pagina de Proyectos puede quedar mucho mejor si cada caso recibe mas imagenes y, si existe, un video o clips de apoyo.",
            "Tambien hace falta una imagen social/OG para compartir la web correctamente en WhatsApp y redes.",
        ],
    )

    add_section(
        document,
        "2. Assets que hay que pedirle al cliente",
        [
            "Logo en SVG o vector editable.",
            "Logo en PNG con fondo transparente.",
            "Version horizontal del logo y version isotipo, si existen.",
            "Favicon o logo cuadrado en alta calidad.",
            "Manual de marca, si existe.",
            "1 foto principal del equipo para la seccion Studio, horizontal y en alta resolucion.",
            "2 a 4 fotos extra del equipo, oficina o behind the scenes para reforzar la seccion Studio.",
            "4 clips hero para la home, uno por cada caso destacado.",
            "1 poster o frame principal por cada clip hero.",
            "Para cada proyecto del portfolio: 1 portada principal y 4 a 8 imagenes extra.",
            "Para cada proyecto del portfolio: 1 video master, video corto o 1 a 3 clips de apoyo, si existen.",
            "Para colorizacion: 3 a 6 trabajos que si quieran mostrar.",
            "Para cada trabajo de colorizacion: 3 a 5 stills finales ya colorizados.",
            "Idealmente, para colorizacion: version raw de esos mismos stills o frames equivalentes.",
            "1 imagen OG/social para compartir la web en WhatsApp, Instagram o enlaces.",
        ],
    )

    add_section(
        document,
        "3. Pedido detallado por seccion",
        [
            "Branding: logo vectorial, PNG transparente, versiones alternativas y cualquier guia visual disponible.",
            "Studio: una foto nueva, grande, horizontal y profesional del equipo. Debe verse premium y servir bien en desktop sin pixelarse.",
            "Studio: si tienen fotos del equipo trabajando, rodajes, set, camaras o proceso, enviarlas tambien.",
            "Home: 4 clips limpios, sin textos quemados, de aproximadamente 6 a 12 segundos cada uno.",
            "Home: enviar tambien un frame fijo o poster por cada clip, por si se necesita fallback en movil o carga lenta.",
            "Proyectos: para La Bodega Dia de los Padres, Shibuya Casa de Campo, Changan Dominicana, Farma Extra y Porsche Center Santo Domingo, enviar portada principal y galeria.",
            "Proyectos: incluir para cada caso el mejor material disponible, priorizando imagenes finales de calidad comercial.",
            "Proyectos: si existe video, enviar el master o cortes breves listos para web.",
            "Colorizacion: como la seccion se simplificara, enviar stills seleccionados en vez de depender del slider comparativo.",
            "Colorizacion: elegir trabajos donde el color realmente se note y represente bien su nivel.",
        ],
    )

    add_numbered_section(
        document,
        "4. Pedido por proyecto",
        [
            "La Bodega Dia de los Padres: 1 portada principal, 4 a 8 imagenes extra, 1 video o clips, y nota corta del caso.",
            "Shibuya Casa de Campo: 1 portada principal, 4 a 8 imagenes extra, 1 video o clips, y nota corta del caso.",
            "Changan Dominicana: 1 portada principal, 4 a 8 imagenes extra, 1 video o clips, y nota corta del caso.",
            "Farma Extra: 1 portada principal, 4 a 8 imagenes extra, 1 video o clips, y nota corta del caso.",
            "Porsche Center Santo Domingo: 1 portada principal, 4 a 8 imagenes extra, 1 video o clips, y nota corta del caso.",
        ],
    )

    add_section(
        document,
        "5. Textos e informacion complementaria que tambien conviene pedir",
        [
            "Texto corto sobre el estudio: quienes son, como trabajan y que tipo de clientes atienden.",
            "Si quieren, nombres y roles del equipo principal.",
            "Breve descripcion por proyecto: cliente, ano, objetivo, entregables y que hizo Manzana Cuatro.",
            "Si tienen testimonios o frases de clientes, tambien conviene pedirlas.",
        ],
    )

    add_section(
        document,
        "6. Aclaracion sobre 'stills'",
        [
            "Stills se escribe asi: stills.",
            "En audiovisual, un still es una imagen fija tomada de un video o de una pieza filmada.",
            "Tambien se puede explicar al cliente como: frames, fotogramas o capturas finales del trabajo.",
            "Para la seccion de Colorizacion, los stills son suficientes si tienen buena calidad y muestran bien el look final.",
        ],
    )

    add_section(
        document,
        "7. Recomendaciones de entrega para el cliente",
        [
            "Imagenes en JPG o PNG de alta calidad, idealmente minimo 2500 a 3000 px de ancho.",
            "Videos en MP4, preferiblemente H.264.",
            "Evitar archivos descargados de WhatsApp o capturas comprimidas.",
            "Si pueden, enviar el material ordenado por carpetas: Home, Studio, Proyectos y Colorizacion.",
            "Nombrar los archivos con claridad para evitar confusion al montar la web.",
        ],
    )

    closing = document.add_paragraph()
    closing.add_run("Nota final: ").bold = True
    closing.add_run(
        "Con este material se puede cerrar la web con una presentacion mucho mas fuerte en Home, "
        "Studio, Proyectos y Colorizacion."
    )

    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
